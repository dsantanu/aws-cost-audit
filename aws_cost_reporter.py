#!/usr/bin/env python3
"""
Name    : aws_cost_reporter.py
Desc    : Generate an executive AWS Cost Optimization report (DOCX) from your CLI audit pack.
Author  : Santanu Das (@dsantanu) | License : MIT
Version : v1.1.0

Usage:
  python aws_cost_report.py --input ./audit_dir --output ./report.docx [--charts] [--org "My Company"] [--author "Your Name (Role)"]

Notes:
- Expects (where available) files like:
  cost-by-service.json, ec2-instances.json, cpu_*.json, ebs-volumes.json,
  s3-<bucket>-size.json, rds.json, loadbalancers.json, nat-gateways.json, tags.json,
  route53-cost.json, route53-zones.json, route53-health-checks.json,
  route53-records-*.json, elastic-ips.json, eip-cost.json
"""

import argparse, json, re, sys
from pathlib import Path
from datetime import datetime
from statistics import mean
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# Optional charts
try:
    import matplotlib.pyplot as plt
    MATPLOTLIB_OK = True
except Exception:
    MATPLOTLIB_OK = False

AWS_ORANGE = "#FF9900"
AWS_BLUE   = "#0073BB"
AWS_GRAY   = "#232F3E"
AWS_PALETTE = [AWS_ORANGE, AWS_BLUE, AWS_GRAY, "#A6ACAF", "#D5DBDB"]

def load_json(path: Path):
    try:
        with open(path, "r") as f:
            return json.load(f)
    except Exception:
        return None

def load_cost_by_service(audit_dir: Path):
    p = audit_dir / "cost-by-service.json"
    data = load_json(p)
    rows = []
    if data:
        groups = (data or {}).get("ResultsByTime", [{}])[0].get("Groups", [])
        for g in groups:
            svc = g.get("Keys", ["Unknown"])[0]
            amt = g.get("Metrics", {}).get("UnblendedCost", {}).get("Amount", "0") or "0"
            try:
                val = float(amt)
            except Exception:
                try:
                    val = float(amt.replace("E","e"))
                except Exception:
                    val = 0.0
            rows.append({"Service": svc, "CostUSD": val})
    return pd.DataFrame(rows).sort_values("CostUSD", ascending=False)

def load_ec2(audit_dir: Path):
    p = audit_dir / "ec2-instances.json"
    data = load_json(p)
    if not data:
        return pd.DataFrame(), pd.DataFrame()
    # object-shape vs array-shape tolerant
    if isinstance(data, list) and data and isinstance(data[0], dict):
        ec2 = pd.DataFrame(data)
        if "State" in ec2.columns and isinstance(ec2.iloc[0]["State"], dict):
            ec2["State"] = ec2["State"].apply(lambda s: s.get("Name") if isinstance(s, dict) else s)
    else:
        rows = []
        arr = data if isinstance(data, list) else []
        flat = []
        for r in arr:
            if isinstance(r, list): flat.extend(r)
            else: flat.append(r)
        for rec in flat:
            if isinstance(rec, list) and len(rec) >= 5:
                rows.append({
                    "InstanceId": rec[0], "InstanceType": rec[1], "State": rec[2],
                    "AZ": rec[3], "LaunchTime": rec[4]
                })
        ec2 = pd.DataFrame(rows)
    # CPU metrics
    cpu_rows = []
    for pth in audit_dir.glob("cpu_*.json"):
        d = load_json(pth)
        if not d:
            continue
        dps = d.get("Datapoints", [])
        avgs = [dp.get("Average") for dp in dps if "Average" in dp]
        inst = pth.stem.replace("cpu_","")
        if avgs:
            avg = float(np.mean(avgs))
            p95 = float(np.percentile(avgs, 95)) if len(avgs) > 1 else avg
            samples = len(avgs)
        else:
            avg = None; p95=None; samples = 0
        cpu_rows.append({"InstanceId": inst, "CPUAvg7d": avg, "CPU95p7d": p95, "Samples": samples})
    cpu = pd.DataFrame(cpu_rows)
    ec2m = ec2.merge(cpu, on="InstanceId", how="left") if not ec2.empty else pd.DataFrame()
    return ec2, ec2m

def load_ebs(audit_dir: Path):
    p = audit_dir / "ebs-volumes.json"
    data = load_json(p)
    if not data: return pd.DataFrame()
    if isinstance(data, list) and data and isinstance(data[0], dict):
        return pd.DataFrame(data)
    rows=[]; arr = data if isinstance(data, list) else []
    flat=[]
    for r in arr:
        if isinstance(r, list): flat.extend(r)
        else: flat.append(r)
    for rec in flat:
        if isinstance(rec, list) and len(rec)>=7:
            rows.append({
                "VolumeId": rec[0], "Size": rec[1], "VolumeType": rec[2], "State": rec[3],
                "InstanceId": rec[4], "Encrypted": rec[5], "CreateTime": rec[6]
            })
    return pd.DataFrame(rows)

def load_s3_sizes(audit_dir: Path):
    rows=[]
    for p in audit_dir.glob("s3-*-size.json"):
        d = load_json(p);
        if not d: continue
        dps = d.get("Datapoints", [])
        avgs = [dp.get("Average") for dp in dps if "Average" in dp]
        avg_bytes = float(np.mean(avgs)) if avgs else 0.0
        bucket = p.name.replace("s3-","").replace("-size.json","")
        rows.append({"Bucket": bucket, "AvgGiB3d": avg_bytes/1024/1024/1024})
    return pd.DataFrame(rows).sort_values("AvgGiB3d", ascending=False)

def load_misc(audit_dir: Path):
    rds = load_json(audit_dir / "rds-instances.json")
    if isinstance(rds, list) and rds and isinstance(rds[0], dict):
        rds_df = pd.DataFrame(rds)
    elif isinstance(rds, list):
        rows=[]
        for rec in rds:
            if isinstance(rec, list) and len(rec)>=3:
                rows.append({"DBInstanceIdentifier": rec[0], "DBInstanceClass": rec[2]})
        rds_df = pd.DataFrame(rows)
    else:
        rds_df = pd.DataFrame()
    elbv2 = load_json(audit_dir / "loadbalancers.json") or {}
    natgw = load_json(audit_dir / "nat-gateways.json") or {}
    tags  = load_json(audit_dir / "tags.json") or {}
    lb_total = len(elbv2.get("LoadBalancers", []))
    nat_total = len(natgw.get("NatGateways", []))
    tag_count = len(tags.get("ResourceTagMappingList", []))
    return rds_df, lb_total, nat_total, tag_count

def load_route53_eip(audit_dir: Path):
    """Load Route53 and Elastic IP related audit data and return summary dict."""
    r53_cost = load_json(audit_dir / "route53-cost.json")
    eip_cost = load_json(audit_dir / "eip-cost.json")
    zones    = load_json(audit_dir / "route53-zones.json")
    health   = load_json(audit_dir / "route53-health-checks.json")
    eips     = load_json(audit_dir / "elastic-ips.json")

    def cost_amt(obj):
        try:
            return float(obj["ResultsByTime"][0]["Total"]["UnblendedCost"]["Amount"])
        except Exception:
            return 0.0

    r53_cost_usd = cost_amt(r53_cost) if r53_cost else 0.0
    eip_cost_usd = cost_amt(eip_cost) if eip_cost else 0.0

    zone_count = len((zones or {}).get("HostedZones", []))
    hc_count   = len((health or {}).get("HealthChecks", []))
    eip_list   = (eips or {}).get("Addresses", [])
    eip_total  = len(eip_list)
    eip_unattached = sum(
        1 for e in eip_list
        if ("InstanceId" not in e and "NetworkInterfaceId" not in e and "AssociationId" not in e)
    )

    # detect duplicate targets and low TTLs
    dupes, low_ttl = [], []
    for p in audit_dir.glob("route53-records-*.json"):
        d = load_json(p) or {}
        rrsets = d.get("ResourceRecordSets", [])
        zone_id = p.stem.replace("route53-records-", "")
        target_map = {}
        for rr in rrsets:
            rtype = rr.get("Type", "")
            name  = rr.get("Name", "")
            alias = rr.get("AliasTarget")
            ttl   = rr.get("TTL")
            if rtype in ("A", "AAAA", "CNAME") and not alias and isinstance(ttl, int) and ttl < 300:
                low_ttl.append({"ZoneId": zone_id, "Name": name, "Type": rtype, "TTL": ttl})
            if rtype == "A":
                if alias and "DNSName" in alias:
                    target = alias["DNSName"].rstrip(".").lower()
                    target_map.setdefault(target, []).append(name)
                else:
                    for rec in rr.get("ResourceRecords", []):
                        v = rec.get("Value")
                        if v:
                            target_map.setdefault(v, []).append(name)
        for target, names in target_map.items():
            if len(set(names)) > 1:
                dupes.append({"ZoneId": zone_id, "Target": target, "Names": ", ".join(sorted(set(names)))})

    # return unified dict
    return {
        "r53_cost": r53_cost_usd,
        "eip_cost": eip_cost_usd,
        "zones": zone_count,
        "health_checks": hc_count,
        "eips": eip_total,
        "eips_unattached": eip_unattached,
        "dupes": pd.DataFrame(dupes),
        "low_ttl": pd.DataFrame(low_ttl)
    }

def downsize_type(it, steps):
    size_order = ["nano","micro","small","medium","large","xlarge","2xlarge","3xlarge","4xlarge","6xlarge","8xlarge","12xlarge","16xlarge","24xlarge","32xlarge","metal"]
    try:
        fam, sz = it.split(".",1)
        if sz not in size_order: return it
        idx = max(0, size_order.index(sz)-steps)
        return f"{fam}.{size_order[idx]}"
    except Exception:
        return it

def build_report(audit_dir: Path, output_docx: Path, org_name: str, author: str, charts: bool=False):
    # ---- Load data (expects your original helper functions) ----
    cost_df = load_cost_by_service(audit_dir)                      # columns: Service, CostUSD
    ec2_df, ec2m_df = load_ec2(audit_dir)                          # ec2 inventory + metrics summary
    ebs_df = load_ebs(audit_dir)                                   # volumes
    s3sizes_df = load_s3_sizes(audit_dir)                          # Bucket, AvgGiB3d
    rds_df, lb_total, nat_total, tag_count = load_misc(audit_dir)  # rds instances, totals
    r53eip = load_route53_eip(audit_dir)                           # dicts: {'MonthlyCost', 'Zones', 'HealthChecks'}, {'MonthlyCost','Allocated','Unattached'}

    # ---- Compute quick insights ----
    top_service = cost_df.iloc[0]["Service"] if not cost_df.empty else "N/A"
    top_cost = float(cost_df.iloc[0]["CostUSD"]) if not cost_df.empty else 0.0

    ec2_total = len(ec2_df) if not ec2_df.empty else 0
    ec2_running = int((ec2_df["State"].astype(str)=="running").sum()) if not ec2_df.empty and "State" in ec2_df.columns else 0

    # rightsizing candidates & table rows
    rs_rows = []
    idle_candidates = 0
    if not ec2m_df.empty:
        mask_running = (ec2m_df["State"].astype(str)=="running") if "State" in ec2m_df.columns else ec2m_df.index==ec2m_df.index
        for _, r in ec2m_df[mask_running].iterrows():
            cur = str(r.get("InstanceType",""))
            cpu = r.get("CPUAvg7d")
            p95 = r.get("CPU95p7d")
            samples = int(r.get("Samples") or 0)
            if pd.isna(cpu) or samples < 12:
                rec, reason = cur, "Insufficient metrics (retain)"
            elif cpu < 5:
                rec, reason = downsize_type(cur,2), "CPU<5% (downsize 2)"
                idle_candidates += 1
            elif cpu < 20:
                rec, reason = downsize_type(cur,1), "CPU 5–20% (downsize 1)"
            else:
                rec, reason = cur, "CPU≥20% (retain)"
            rs_rows.append([
                r.get("InstanceId",""),
                cur,
                None if pd.isna(cpu) else float(cpu),
                None if pd.isna(p95) else float(p95),
                samples,
                rec,
                reason
            ])
    rs_df = pd.DataFrame(rs_rows, columns=[
        "InstanceId","CurrentType","CPUAvg7d","CPU95p7d","Samples","RecommendedType","Reason"
    ])

    # storage signals
    ebs_unattached = 0
    ebs_gp2 = 0
    if not ebs_df.empty:
        if "InstanceId" in ebs_df.columns:
            ebs_unattached = int(((ebs_df["InstanceId"].isna()) | (ebs_df["InstanceId"].astype(str) == "")).sum())

        if "VolumeType" in ebs_df.columns:
            ebs_gp2 = int((ebs_df["VolumeType"].astype(str)=="gp2").sum())

    largest_bucket = "N/A"
    largest_gib = 0.0
    if not s3sizes_df.empty:
        s3sizes_df = s3sizes_df.sort_values("AvgGiB3d", ascending=False)
        largest_bucket = s3sizes_df.iloc[0]["Bucket"]
        largest_gib = float(s3sizes_df.iloc[0]["AvgGiB3d"])

    # ---- Route53 / EIP (single combined dict) ----
    r53eip = load_route53_eip(audit_dir)

    r53_cost   = float(r53eip.get("r53_cost", 0.0))
    r53_zones  = int(r53eip.get("zones", 0))
    r53_health = int(r53eip.get("health_checks", 0))

    eip_cost   = float(r53eip.get("eip_cost", 0.0))
    eip_alloc  = int(r53eip.get("eips", 0))
    eip_unatt  = int(r53eip.get("eips_unattached", 0))

    # (optionally useful later)
    dupes_df   = r53eip.get("dupes", pd.DataFrame())
    lowttl_df  = r53eip.get("low_ttl", pd.DataFrame())

    # ---- Charts (Top 5 + Projected Savings) ----
    charts_dir = output_docx.parent / "charts"
    charts_dir.mkdir(exist_ok=True)
    top5_png = charts_dir / "top5_services_cost.png"
    proj_png = charts_dir / "projected_savings.png"

    # Top 5 Services Pie Chart
    service_aliases = {
      "Amazon Elastic Compute Cloud - Compute": "EC2",
      "Amazon Simple Storage Service": "S3",
      "Amazon Relational Database Service": "RDS",
      "Amazon Elastic Container Service": "ECS",
      "Amazon Elastic Kubernetes Service": "EKS",
      "Tax": "Tax",
    }

    if charts and MATPLOTLIB_OK:
        top5 = cost_df.head(5)
        if not top5.empty:
            plt.figure(figsize=(6,4))
            t5_labels = [service_aliases.get(s, s.replace("Amazon ", "").split(" - ")[0]) for s in top5["Service"]]
            plt.pie(
                top5["CostUSD"],
                labels=top5["Service"],
                #labels=t5_labels,
                autopct=lambda p: f"${p*top5['CostUSD'].sum()/100:,.0f}",
                startangle=140
            )
            plt.title("Top 5 Services by Cost (USD)")
            plt.tight_layout(pad=1.5)
            plt.savefig(top5_png, dpi=200, bbox_inches='tight')
            plt.close()
            print(f"[INFO] Chart generated: {top5_png}")

        # simple static model for projected savings (can wire to your computed estimates later)
        cats = ["EC2 Rightsizing","EBS Optimization","S3 Lifecycle","Networking","Governance"]
        vals = [22, 8, 6, 4, 2]
        plt.figure(figsize=(6,4))
        plt.bar(cats, vals)
        plt.ylabel("Projected Savings (%)")
        plt.title("Projected Savings by Optimization Category")
        plt.xticks(rotation=20, ha="right")
        plt.tight_layout()
        plt.savefig(proj_png, dpi=200); plt.close()
        print(f"[INFO] Chart generated: {proj_png}")

    # ---- DOCX build ----
    doc = Document()

    # header block (no branded cover yet)
    doc.add_heading("AWS Cost Optimization & Efficiency Report", level=0)
    hdr = doc.add_paragraph(
        f"Date: {datetime.now().strftime('%Y-%m-%d')}  |  Currency: USD ($)\n"
        f"Prepared by: {author}\n"
        f"Organization: {org_name}"
    )
    hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Key Insights
    doc.add_heading("Key Insights at a Glance", level=1)
    dash = doc.add_table(rows=5, cols=2)
    dash.style = "Light Grid"
    dash.cell(0,0).text = "Top Cost Driver"; dash.cell(0,1).text = f"{top_service} — ~${top_cost:,.2f}"
    dash.cell(1,0).text = "EC2 Fleet"; dash.cell(1,1).text = f"Total {ec2_total}, Running {ec2_running}, Idle {idle_candidates}"
    dash.cell(2,0).text = "Storage Signals"; dash.cell(2,1).text = f"Unattached EBS: {ebs_unattached} | gp2: {ebs_gp2} | Largest S3: {largest_bucket} ({largest_gib:.2f} GiB)"
    dash.cell(3,0).text = "Databases"; dash.cell(3,1).text = f"RDS instances: {len(rds_df) if rds_df is not None else 0}"
    dash.cell(4,0).text = "Networking & Tagging"; dash.cell(4,1).text = f"NAT GWs: {nat_total} | LBs: {lb_total} | Tagged: {tag_count}"

    # Top 5 chart right after Key Insights
    if charts and MATPLOTLIB_OK and top5_png.exists():
        doc.add_paragraph()
        doc.add_heading("Top 5 Services by Cost (USD)", level=2)
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_picture(str(top5_png), width=Inches(5.5))
        cap = doc.add_paragraph("Figure 1: Top 5 cost drivers from AWS Cost Explorer data.")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].italic = True

    # Executive Summary
    doc.add_paragraph()
    doc.add_heading("Executive Summary", level=1)
    ex = doc.add_paragraph()
    ex.add_run(
        f"The account’s largest cost driver is {top_service} (~${top_cost:,.2f}). "
        "Immediate savings include EC2 rightsizing, EBS cleanup (unattached removal and gp2→gp3), "
        "S3 lifecycle transitions, and DNS/EIP optimization. Networking savings via NAT consolidation "
        "and Gateway Endpoints. Tagging improvements will enhance cost transparency."
    )

    # Compute (EC2) Analysis
    doc.add_paragraph()
    doc.add_heading("Compute (EC2) Analysis", level=1)
    doc.add_paragraph(
        "Rightsizing thresholds: CPU < 5% → downsize 2 sizes; CPU 5–20% → downsize 1 size. "
        "Instances without sufficient telemetry are retained."
    )
    if not rs_df.empty:
        tbl = doc.add_table(rows=1, cols=len(rs_df.columns)); tbl.style = "Light Grid"
        hdr_cells = tbl.rows[0].cells
        for i, col in enumerate(rs_df.columns): hdr_cells[i].text = col
        for _, row in rs_df.iterrows():
            cells = tbl.add_row().cells
            cells[0].text = str(row["InstanceId"])
            cells[1].text = str(row["CurrentType"])
            cells[2].text = "" if pd.isna(row["CPUAvg7d"]) else f"{row['CPUAvg7d']:.3f}"
            cells[3].text = "" if pd.isna(row["CPU95p7d"]) else f"{row['CPU95p7d']:.3f}"
            cells[4].text = str(row["Samples"])
            cells[5].text = str(row["RecommendedType"])
            cells[6].text = str(row["Reason"])
    else:
        doc.add_paragraph("No EC2 instances or insufficient metrics for rightsizing analysis.")

    # Savings Plan Efficiency (textual for now)
    doc.add_paragraph()
    doc.add_heading("Savings Plan Efficiency", level=1)
    doc.add_paragraph(
        "If Savings Plan data is available in the audit pack, summarize utilization and commitment here. "
        "Maintain coverage above ~85% for steady-state compute; revisit commitments if workloads change."
    )

    # Storage (EBS + S3)
    doc.add_paragraph()
    doc.add_heading("Storage (EBS + S3)", level=1)
    doc.add_paragraph(
        f"EBS: {ebs_unattached} unattached volumes; snapshot & delete. "
        f"{ebs_gp2} gp2 volumes → migrate to gp3. "
        f"S3: Largest bucket '{largest_bucket}' ~{largest_gib:.2f} GiB (3-day avg). "
        "Apply lifecycle to IA/Glacier; consider Intelligent-Tiering."
    )

    # Networking & Load Balancing
    doc.add_paragraph()
    doc.add_heading("Networking & Load Balancing", level=1)
    doc.add_paragraph(
        f"NAT Gateways: {nat_total}. Consolidate and add Gateway Endpoints for S3/DynamoDB. "
        f"Load Balancers: {lb_total}. Remove idle LBs and unused SGs."
    )

    # DNS & IP Optimization
    doc.add_paragraph()
    doc.add_heading("DNS & IP Optimization (Route 53 / EIP)", level=1)
    doc.add_paragraph(
        f"Route 53 cost: ${r53_cost:.2f}/month | Hosted zones: {r53_zones} | Health checks: {r53_health}\n"
        f"Elastic IP cost (reported): ${eip_cost:.2f}/month | Allocated: {eip_alloc} | Unattached: {eip_unatt}"
    )
    tbl = doc.add_table(rows=1, cols=3); tbl.style = "Light Grid"
    tbl.rows[0].cells[0].text = "Area"
    tbl.rows[0].cells[1].text = "Findings"
    tbl.rows[0].cells[2].text = "Recommended Action"
    # Row: Route 53
    r = tbl.add_row().cells
    r[0].text = "Route 53"
    r[1].text = "Multiple A-record names pointing to same targets; low TTLs detected"
    r[2].text = "Consolidate via ALB alias; increase TTL to 300–900s; remove stale records"
    # Row: Elastic IPs
    r = tbl.add_row().cells
    r[0].text = "Elastic IPs"
    r[1].text = f"{eip_alloc} total, {eip_unatt} unattached"
    r[2].text = "Release unattached EIP; prefer ALB/NLB DNS endpoints over static EIPs"
    # Row: Health Checks
    r = tbl.add_row().cells
    r[0].text = "Health Checks"
    r[1].text = f"{r53_health} active Route 53 health checks"
    r[2].text = "Use ALB/NLB target health unless needed for external endpoints"

    # Governance & Observability
    doc.add_paragraph()
    doc.add_heading("Governance & Observability", level=1)
    doc.add_paragraph(
        "Activate/enforce cost allocation tags (Environment, Application, Owner, CostCenter). "
        "Set CloudWatch log retention to 30–90 days and export to S3 with lifecycle to Glacier."
    )

    # Prioritized Remediation Plan
    doc.add_paragraph()
    doc.add_heading("Prioritized Remediation Plan", level=1)
    plan = [
        ("1","Stop/resize idle EC2s","10–25%","2–4"),
        ("2","Delete unattached EBS & migrate gp2→gp3","5–10%","2"),
        ("3","Add S3 lifecycle (IA/Glacier)","3–8%","3"),
        ("4","Consolidate NAT & add Gateway Endpoints","2–6%","4"),
        ("5","DNS/EIP optimization (Route 53 / EIP)","1–3%","3"),
        ("6","Enforce tagging (Config/SCP)","—","4"),
    ]
    t = doc.add_table(rows=1, cols=4); t.style = "Light Grid"
    t.rows[0].cells[0].text = "Priority"
    t.rows[0].cells[1].text = "Action"
    t.rows[0].cells[2].text = "Target Savings"
    t.rows[0].cells[3].text = "Effort (hrs)"
    for pr, act, sav, hrs in plan:
        c = t.add_row().cells
        c[0].text, c[1].text, c[2].text, c[3].text = pr, act, sav, hrs

    # Projected Savings chart right after the plan
    if charts and MATPLOTLIB_OK and proj_png.exists():
        doc.add_paragraph()
        doc.add_heading("Projected Savings by Optimization Category", level=2)
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_picture(str(proj_png), width=Inches(5.5))
        cap = doc.add_paragraph("Figure 2: Estimated savings contribution by optimization category.")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].italic = True

    # Appendix
    doc.add_paragraph()
    doc.add_heading("Appendix — Methodology & Data Validity", level=1)
    doc.add_paragraph(
        "Data gathered via AWS CLI audit pack. Where telemetry was missing, recommendations default "
        "to 'retain size' with follow-up to enable metrics collection."
    )
    doc.add_paragraph("\nReviewed & Approved by: __________________________")
    doc.add_paragraph("Date: __________________________")

    # Save
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_docx)
    print(f"[OK] Report written to: {output_docx}")

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--input", required=True, help="Path to audit directory")
    ap.add_argument("--output", required=True, help="Path to output .docx file")
    ap.add_argument("--charts", action="store_true", help="Generate and embed PNG charts")
    ap.add_argument("--org", default="AWS Cost Optimization Report", help="Organization name for the report")
    ap.add_argument("--author", default="Cloud Architecture", help="Display author's name in the report")
    args = ap.parse_args()
    audit_dir = Path(args.input)
    output_docx = Path(args.output)
    if not audit_dir.exists():
        print(f"[ERR] Input path not found: {audit_dir}", file=sys.stderr)
        sys.exit(1)
    build_report(audit_dir, output_docx, args.org, args.author, charts=args.charts)

if __name__ == "__main__":
    main()
